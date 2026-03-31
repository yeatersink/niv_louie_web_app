# utils/braille_document_manager.py

from nicegui import events, ui
import io
import os
from pathlib import Path
import json
from docx import Document
import sys
import pandas as pd

from utils.storage import get_user_documents_dir, ensure_user_directories, get_user_projects_dir
from utils.braille import get_braille_from_text
from utils.project import project

# Load braille conversion files
if getattr(sys, 'frozen', False):
    base_path = sys._MEIPASS
else:
    base_path = os.path.abspath(".")

try:
    with open(os.path.join(base_path, "utils", "braille_to_numbers.json"), encoding="utf8") as f:
        braille_numbers_object = json.load(f)
    with open(os.path.join(base_path, "utils", "braille_test_converter.json"), encoding="utf8") as f:
        braille_converter_object = json.load(f)
except Exception as ex:
    print(f"ERROR loading braille JSON files: {ex}")
    braille_numbers_object = {}
    braille_converter_object = {}


class DocumentManager:
    def __init__(self):
        self.document_name = None
        self.document_contents = None
        self.selected_projects = []          # list for multiple projects, order matters
        self.apply_general_english = True    # toggle for general English rules
        self.document_text = None            # temporary storage for conversion

    def update_document_name(self, e: events.ValueChangeEventArguments):
        self.document_name = e.value
        ui.notify(f"Selected document: {e.value}")

    def update_selected_projects(self, e: events.ValueChangeEventArguments):
        self.selected_projects = e.value if e.value else []
        if self.selected_projects:
            ui.notify(f"Selected projects: {', '.join(self.selected_projects)}")
        else:
            ui.notify("No projects selected", type="warning")

    def toggle_general_english(self, e: events.ValueChangeEventArguments):
        self.apply_general_english = e.value
        status = "enabled" if self.apply_general_english else "disabled"
        ui.notify(f"General English rules {status}")

    async def handle_document_upload(self, e: events.UploadEventArguments):
        document_folder = get_user_documents_dir()
        document_folder.mkdir(parents=True, exist_ok=True)

        try:
            if not e.file or not e.file.name:
                ui.notify("Upload failed - no file received", type="negative")
                return

            content_bytes = await e.content.read() if hasattr(e, 'content') and e.content else await e.file.read()

            filename = e.file.name.lower()

            if filename.endswith(".docx"):
                document = Document(io.BytesIO(content_bytes))
                save_path = document_folder / e.file.name
                document.save(save_path)
                self.document_name = e.file.name
                self.document_contents = document
                ui.notify(f"✅ Docx document uploaded: {e.file.name}")

            elif filename.endswith(".txt"):
                content_str = content_bytes.decode("utf-8", errors="replace")
                save_path = document_folder / e.file.name
                with open(save_path, "w", encoding="utf-8") as file:
                    file.write(content_str)
                self.document_name = e.file.name
                ui.notify(f"✅ Text document uploaded: {e.file.name}")

            elif filename.endswith(".pdf"):
                import fitz
                doc = fitz.open(stream=content_bytes, filetype="pdf")
                text = ""
                for page in doc:
                    text += page.get_text("text") + "\n"
                doc.close()

                save_path = document_folder / e.file.name
                with open(save_path, "w", encoding="utf-8") as file:
                    file.write(text)
                self.document_name = e.file.name
                ui.notify(f"✅ PDF document uploaded and text extracted: {e.file.name}")

            else:
                ui.notify("Only .docx, .txt, and .pdf files are supported", type="negative")
                return

        except Exception as ex:
            ui.notify(f"Document upload error: {str(ex)}", type="negative")
            print(f"DEBUG document upload error: {ex}")

    def convert_document(self):
        if not self.document_name:
            ui.notify("No document selected", type="negative")
            return

        document_folder = get_user_documents_dir()
        document_path = document_folder / self.document_name

        try:
            # Extract plain text from the uploaded document
            if self.document_name.lower().endswith(".docx"):
                document = Document(document_path)
                text = "\n".join([paragraph.text for paragraph in document.paragraphs])
            else:
                with open(document_path, "r", encoding="utf-8") as f:
                    text = f.read()

            self.document_text = text

            # Show dialog to choose output format
            with ui.dialog() as dialog, ui.card().classes('w-96'):
                ui.label("Choose output format for the braille document:").classes('text-lg')

                def download_as(fmt: str):
                    braille_content = self.convert_text_to_braille(self.document_text)
                    base_name = self.document_name.rsplit('.', 1)[0]

                    if fmt == "txt":
                        output_filename = f"{base_name}_braille.txt"
                        output_path = document_folder / output_filename
                        with open(output_path, "w", encoding="utf-8") as f:
                            f.write(braille_content)
                        
                        # Fixed: Use ui.download.content instead of ui.download(path)
                        with open(output_path, "rb") as f:
                            content = f.read()
                        ui.download.content(content, filename=output_filename, media_type="text/plain")

                    elif fmt == "brf":
                        output_filename = f"{base_name}_braille.brf"
                        output_path = document_folder / output_filename
                        with open(output_path, "w", encoding="utf-8") as f:
                            f.write(braille_content)
                        
                        with open(output_path, "rb") as f:
                            content = f.read()
                        ui.download.content(content, filename=output_filename, media_type="text/plain")

                    elif fmt == "docx":
                        output_filename = f"{base_name}_braille.docx"
                        output_path = document_folder / output_filename
                        doc = Document()
                        doc.add_paragraph(braille_content)
                        doc.save(output_path)
                        
                        with open(output_path, "rb") as f:
                            content = f.read()
                        ui.download.content(content, filename=output_filename, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

                    dialog.close()
                    ui.notify(f"Downloaded as .{fmt}")

                with ui.column().classes('w-full gap-4'):
                    ui.button("Download as .txt (Plain Text)", on_click=lambda: download_as("txt")).classes('w-full')
                    ui.button("Download as .brf (Braille Ready Format)", on_click=lambda: download_as("brf")).classes('w-full')
                    ui.button("Download as .docx (Word Document)", on_click=lambda: download_as("docx")).classes('w-full')

            dialog.open()

        except Exception as ex:
            ui.notify(f"Error converting document: {str(ex)}", type="negative")

    def remove_document(self):
        if not self.document_name:
            ui.notify("No document selected", type="negative")
            return

        document_folder = get_user_documents_dir()
        document_path = document_folder / self.document_name

        try:
            if document_path.exists():
                document_path.unlink()
                ui.notify(f"Document {self.document_name} removed")
                self.document_name = None
            else:
                ui.notify("Document not found", type="negative")
        except Exception as ex:
            ui.notify(f"Error removing document: {str(ex)}", type="negative")

    def convert_text_to_braille(self, text: str) -> str:
        """Apply selected projects in order, then optional general English rules"""
        braille_content = ""
        for original_line in text.split("\n"):
            current = original_line.strip()

            # Apply each selected project in order
            for proj_name in self.selected_projects:
                try:
                    project.set_project_name(proj_name)
                    project.set_all_fields()

                    projects_dir = get_user_projects_dir()
                    filtered_path = projects_dir / f"filtered_{proj_name}.csv"

                    if filtered_path.exists():
                        df = pd.read_csv(filtered_path)
                        char_col = project.project_character_column
                        braille_col = project.project_braille_column

                        char_map = dict(zip(df[char_col], df[braille_col]))
                        for char, braille in char_map.items():
                            if str(char) in current and str(braille) != "nan":
                                current = current.replace(str(char), str(braille))
                except Exception as e:
                    print(f"  → ERROR applying project {proj_name}: {e}")

            # Apply general English rules ONLY if toggle is on
            if self.apply_general_english:
                for char in list(current):
                    if char in braille_converter_object:
                        current = current.replace(char, braille_converter_object[char])

            braille_content += current + "\n"

        return braille_content


document = DocumentManager()