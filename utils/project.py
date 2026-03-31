# utils/project.py

from nicegui import app, events, ui
import io
from pathlib import Path
import pandas as pd
import json
from docx import Document
import warnings

from utils.storage import get_user_projects_dir, get_user_documents_dir, get_user_tests_dir, ensure_user_directories


# Load braille conversion files (app resources)
braille_test_file = open("utils/braille_test_converter.json", encoding="utf8")
braille_converter_object = json.load(braille_test_file)

braille_numbers_file = open("utils/braille_to_numbers.json", encoding="utf8")
braille_numbers_object = json.load(braille_numbers_file)


def load_languages():
    """Safe load - creates languages_file.json if it doesn't exist."""
    try:
        ensure_user_directories()
        languages_file = get_user_projects_dir() / "languages_file.json"
        
        print(f"LOG: Attempting to load languages from: {languages_file}")

        if not languages_file.exists():
            languages_file.parent.mkdir(parents=True, exist_ok=True)
            with open(languages_file, "w", encoding="utf-8") as file:
                json.dump([], file, ensure_ascii=False, indent=4)
            print(f"LOG: Created empty languages_file.json for new user")
            return []
        
        with open(languages_file, "r", encoding="utf-8") as file:
            data = json.load(file)
            print(f"LOG: Successfully loaded {len(data)} languages from file")
            return data
    except Exception as ex:
        print(f"LOG: Error loading languages: {ex}")
        import traceback
        traceback.print_exc()
        return []


class Project:
    def __init__(self):
        self.project_name = None
        self.project_text = None
        self.project_name_column = None
        self.project_character_column = None
        self.project_unicode_column = None
        self.project_type_column = None
        self.project_braille_column = None
        self.project_language_code = None
        self.project_language_system_code = None
        self.project_display_name = None
        self.project_index_name = None
        self.project_supported_braille_languages = None
        self.project_language_information = None
        self.project_contributors = None
        self.project_included_braille_tables = None
        self.project_test_display_type = None
        self.project_replace = None
        
        self.languages = []
        self.languages_list = []
        
        self.document_name = None
        self.document_contents = None
        self.document_projects_to_use = None
        
        self.projects_dir = None

    def set_user_storage(self, projects_dir: str = None):
        try:
            if projects_dir is None:
                projects_dir = str(get_user_projects_dir())
            
            self.projects_dir = Path(projects_dir)
            self.projects_dir.mkdir(parents=True, exist_ok=True)
            print(f"LOG: Project now using user storage: {self.projects_dir}")
            self.reload_languages()
        except Exception as ex:
            print(f"LOG: Error in set_user_storage: {ex}")

    def reload_languages(self):
        self.languages = load_languages()
        self.languages_list = [language.get("name", "") for language in self.languages if language.get("name")]
        print(f"LOG: Languages reloaded. Total projects now: {len(self.languages)}")
        if self.languages:
            print(f"LOG: First project name: {self.languages[0].get('name', 'N/A')}")

    def update_document_projects_to_use(self, e: events.ValueChangeEventArguments):
        self.document_projects_to_use = e.value

    def update_languages_list(self):
        self.languages_list = [language.get("name", "") for language in self.languages]

    def load_language_source(self):
        if not self.projects_dir:
            self.projects_dir = get_user_projects_dir()
        source_dir = self.projects_dir / "source"
        source_dir.mkdir(parents=True, exist_ok=True)
        file_path = source_dir / f"{self.project_name}.csv"
        self.project_text = pd.read_csv(file_path)

    def set_project_name(self, project_name):
        self.project_name = project_name

    def update_project_name(self, e: events.ValueChangeEventArguments):
        self.project_name = e.value
        self.set_all_fields()

    def update_project_name_column(self, e: events.ValueChangeEventArguments):
        self.project_name_column = e.value

    def update_project_character_column(self, e: events.ValueChangeEventArguments):
        self.project_character_column = e.value

    def update_project_unicode_column(self, e: events.ValueChangeEventArguments):
        self.project_unicode_column = e.value

    def update_project_type_column(self, e: events.ValueChangeEventArguments):
        self.project_type_column = e.value

    def update_project_braille_column(self, e: events.ValueChangeEventArguments):
        self.project_braille_column = e.value

    def update_project_language_code(self, e: events.ValueChangeEventArguments):
        self.project_language_code = e.value
        
    def update_project_language_system_code(self, e: events.ValueChangeEventArguments):
        self.project_language_system_code = e.value

    def update_project_display_name(self, e: events.ValueChangeEventArguments):
        self.project_display_name = e.value

    def update_project_index_name(self, e: events.ValueChangeEventArguments):
        self.project_index_name = e.value

    def update_project_supported_braille_languages(self, e: events.ValueChangeEventArguments):
        self.project_supported_braille_languages = e.value
        
    def update_project_language_information(self, e: events.ValueChangeEventArguments):
        self.project_language_information = e.value

    def update_project_contributors(self, e: events.ValueChangeEventArguments):
        self.project_contributors = e.value

    def update_project_included_braille_tables(self, e: events.ValueChangeEventArguments):
        self.project_included_braille_tables = e.value

    def update_project_test_display_type(self, e: events.ValueChangeEventArguments):
        self.project_test_display_type = e.value

    def update_project_replace(self, e: events.ValueChangeEventArguments):
        self.project_replace = e.value

    async def handle_file_upload(self, e: events.UploadEventArguments):
        try:
            if not e.file or not e.file.name:
                ui.notify("Upload failed - no file received", type="negative")
                return

            self.project_name = e.file.name.split(".")[0]
            content_bytes = await e.file.read()
            content_as_file = io.StringIO(content_bytes.decode("utf-8"))
            self.project_text = pd.read_csv(content_as_file)
            
            ui.notify(f"✅ File '{e.file.name}' uploaded successfully. Continue to project information.", type="positive")
        except Exception as ex:
            ui.notify(f"Upload error: {str(ex)}", type="negative")
            print(f"DEBUG upload error: {ex}")

    # Fixed - now saves to the exact folder that create_braille_tests expects
    async def handle_test_upload(self, e: events.UploadEventArguments):
        try:
            if not e.file or not e.file.name:
                ui.notify("Upload failed - no file received", type="negative")
                return

            if not self.project_language_code:
                ui.notify("Please select a project from the dropdown FIRST, then upload the test file.", type="negative")
                return

            test_dir = get_user_tests_dir()   # ← Same folder used by create_braille_tests
            test_dir.mkdir(parents=True, exist_ok=True)

            filename = e.file.name.lower()
            content_bytes = await e.file.read()

            # Extract text based on file type
            if filename.endswith(".docx"):
                doc = Document(io.BytesIO(content_bytes))
                text_lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                print(f"DEBUG: Extracted {len(text_lines)} lines from .docx")
            elif filename.endswith(".txt"):
                try:
                    content_str = content_bytes.decode("utf-8")
                except UnicodeDecodeError:
                    content_str = content_bytes.decode("latin1", errors="replace")
                    print("DEBUG: Used latin1 fallback for .txt")
                text_lines = [line.strip() for line in content_str.splitlines() if line.strip()]
                print(f"DEBUG: Extracted {len(text_lines)} lines from .txt")
            elif filename.endswith(".csv"):
                try:
                    content_str = content_bytes.decode("utf-8")
                except UnicodeDecodeError:
                    content_str = content_bytes.decode("latin1", errors="replace")
                    print("DEBUG: Used latin1 fallback for .csv")
                content_as_file = io.StringIO(content_str)
                df = pd.read_csv(content_as_file, on_bad_lines='skip')
                if len(df.columns) >= 1:
                    text_lines = df.iloc[:, 0].astype(str).str.strip().tolist()
                else:
                    text_lines = []
                print(f"DEBUG: Extracted {len(text_lines)} lines from .csv")
            else:
                ui.notify("Only .csv, .txt, and .docx files are supported for tests", type="negative")
                return

            if not text_lines:
                ui.notify("No text content found in the uploaded file", type="negative")
                return

            # Create clean DataFrame with "Text" column
            test_df = pd.DataFrame({"Text": text_lines})
            
            # Save using the EXACT language code the user set in the project
            output_path = test_dir / f"{self.project_language_code}.csv"
            test_df.to_csv(output_path, index=False)
            
            ui.notify(f"✅ Test file saved with {len(text_lines)} sentences for language code '{self.project_language_code}'.")
            print(f"DEBUG test upload successful: saved to {output_path.name}")

        except Exception as ex:
            ui.notify(f"Test upload error: {str(ex)}", type="negative")
            print(f"DEBUG test upload error: {ex}")

    async def handle_document_upload(self, e: events.UploadEventArguments):
        document_folder = get_user_documents_dir()
        document_folder.mkdir(parents=True, exist_ok=True)

        try:
            if not e.file or not e.file.name:
                ui.notify("Upload failed - no file received", type="negative")
                return

            content_bytes = await e.file.read()

            if e.file.name.lower().endswith(".docx"):
                document = Document(io.BytesIO(content_bytes))
                save_path = document_folder / e.file.name
                document.save(save_path)
                self.document_name = e.file.name
                self.document_contents = document
                ui.notify(f"✅ Docx document uploaded: {e.file.name}")
            elif e.file.name.lower().endswith(".txt"):
                content_str = content_bytes.decode("utf-8", errors="replace")
                save_path = document_folder / e.file.name
                with open(save_path, "w", encoding="utf-8") as file:
                    file.write(content_str)
                self.document_name = e.file.name
                ui.notify(f"✅ Text document uploaded: {e.file.name}")
            else:
                ui.notify("Only .docx and .txt files are supported", type="negative")
        except Exception as ex:
            ui.notify(f"Document upload error: {str(ex)}", type="negative")

    def save_project(self):
        error = False
        if self.project_name is None:
            ui.notify("Please enter a name for your project.", type="negative")
            error = True
        if self.project_name_column is None:
            ui.notify("Please select a name column for your project.", type="negative")
            error = True
        if self.project_character_column is None:
            ui.notify("Please select a character column for your project.", type="negative")
            error = True
        if self.project_unicode_column is None:
            ui.notify("Please select a Unicode column for your project.", type="negative")
            error = True
        if self.project_type_column is None:
            ui.notify("Please select a type column for your project.", type="negative")
            error = True
        if self.project_braille_column is None:
            ui.notify("Please select a braille column for your project.", type="negative")
            error = True

        for language in self.languages:
            if self.project_name.lower() == language.get("name", "").lower():
                ui.notify("A project with that name already exists.", type="negative")
                error = True

        if error:
            return

        project_object = {
            "name": self.project_name,
            "name_column": self.project_name_column,
            "char_column": self.project_character_column,
            "braille_column": self.project_braille_column,
            "type_column": self.project_type_column,
            "unicode_column": self.project_unicode_column,
            "language_code": self.project_language_code,
            "language_system_code": self.project_language_system_code,
            "display_name": self.project_display_name,
            "index_name": self.project_index_name,
            "supported_braille_languages": self.project_supported_braille_languages,
            "language_information": self.project_language_information,
            "contributors": self.project_contributors,
            "included_braille_tables": self.project_included_braille_tables,
            "test_display_type": self.project_test_display_type,
            "replace": self.project_replace
        }

        self.languages.append(project_object)
        self.update_languages_list()

        projects_dir = self.projects_dir or get_user_projects_dir()
        projects_dir.mkdir(parents=True, exist_ok=True)

        languages_file = projects_dir / "languages_file.json"
        with open(languages_file, "w", encoding="utf-8") as file:
            json.dump(self.languages, file, ensure_ascii=False, indent=4)

        source_dir = projects_dir / "source"
        source_dir.mkdir(parents=True, exist_ok=True)
        self.project_text.to_csv(source_dir / f"{self.project_name}.csv", index=False)

        ui.navigate.to("/existing_project")
        ui.notify("Project Saved", close_button="Ok")

    def remove_project(self):
        projects_dir = self.projects_dir or get_user_projects_dir()
        languages_file = projects_dir / "languages_file.json"
        
        removed = False
        if self.project_name is not None:
            new_languages = [lang for lang in self.languages if lang.get("name") != self.project_name]
            if len(new_languages) < len(self.languages):
                removed = True
            self.languages = new_languages
            self.update_languages_list()
            
            with open(languages_file, "w", encoding="utf-8") as file:
                json.dump(self.languages, file, ensure_ascii=False, indent=4)
        if removed:
            ui.notify("Project Removed", close_button="Ok")
        else:
            ui.notify("Project not found", close_button="Ok")

    def check_language_names(self, language):
        if language.get("name") == self.project_name:
            return False
        else:
            return True

    def set_all_fields(self):
        if self.project_name is None:
            ui.notify("No Project Selected!", type="negative")
            return
        current_language = None
        for language in self.languages:
            if language.get("name") == self.project_name:
                current_language = language
        if current_language is None:
            ui.notify("Couldn't Find that Project!", type="negative")
            return
        self.project_name_column = current_language.get("name_column")
        self.project_character_column = current_language.get("char_column")
        self.project_unicode_column = current_language.get("unicode_column")
        self.project_type_column = current_language.get("type_column")
        self.project_braille_column = current_language.get("braille_column")
        self.project_language_code = current_language.get("language_code")
        self.project_display_name = current_language.get("display_name")
        if "language_system_code" in current_language:
            self.project_language_system_code = current_language["language_system_code"]
        if "index_name" in current_language:
            self.project_index_name = current_language["index_name"]
        if "supported_braille_languages" in current_language:
            self.project_supported_braille_languages = current_language["supported_braille_languages"]
        self.project_language_information = current_language.get("language_information")
        self.project_contributors = current_language.get("contributors")
        if "included_braille_tables" in current_language:
            self.project_included_braille_tables = current_language["included_braille_tables"]
        if "test_display_type" in current_language:
            self.project_test_display_type = current_language["test_display_type"]
        if "replace" in current_language:
            self.project_replace = current_language["replace"]


# Global instance
project = Project()